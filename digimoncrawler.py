#!/usr/bin/env python3
"""digimoncrawler - Digimon Card Game deck recipe image downloader."""

import os
import sys
import time
import requests
from PIL import Image
from urllib.parse import urlparse, parse_qs
from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENT

BASE_IMG_URL = "https://digimon-cg-guide.com/wp-content/uploads"
CARD_SIZE_MM = (63, 88)
DPI = 300
CARD_SIZE_PX = (int(CARD_SIZE_MM[0] / 25.4 * DPI), int(CARD_SIZE_MM[1] / 25.4 * DPI))
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",
    "Referer": "https://digimon-cg-guide.com/recipe-creater/",
}


def parse_recipe_url(url: str) -> tuple[list[tuple[str, int]], str]:
    parsed = urlparse(url)
    params = parse_qs(parsed.query)

    recipe_str = params.get("recipe", [""])[0]
    deck_name = params.get("deckname", ["deck"])[0]

    if not recipe_str:
        print("Error: 'recipe' parameter not found in URL.")
        sys.exit(1)

    cards = []
    for entry in recipe_str.split("_"):
        # "bt14-0014" → card_id="BT14-001", qty=4 (last digit = quantity)
        qty = int(entry[-1])
        card_id = entry[:-1].upper()
        cards.append((card_id, qty))

    return cards, deck_name


def resize_card_image(filepath: str):
    with Image.open(filepath) as img:
        resized = img.resize(CARD_SIZE_PX, Image.Resampling.LANCZOS)
        resized.save(filepath, dpi=(DPI, DPI))


def download_card_image(card_id: str, save_dir: str) -> bool:
    filename = f"{card_id}.png"
    filepath = os.path.join(save_dir, filename)

    if os.path.exists(filepath):
        print(f"  [SKIP] {filename} (already exists)")
        return True

    url = f"{BASE_IMG_URL}/{card_id}.png"
    try:
        resp = requests.get(url, headers=HEADERS, timeout=30)
        if resp.status_code == 200:
            with open(filepath, "wb") as f:
                f.write(resp.content)
            resize_card_image(filepath)
            size_kb = os.path.getsize(filepath) / 1024
            print(
                f"  [OK]   {filename} ({size_kb:.1f} KB, {CARD_SIZE_PX[0]}x{CARD_SIZE_PX[1]}px)"
            )
            return True
        else:
            print(f"  [FAIL] {filename} (HTTP {resp.status_code})")
            return False
    except requests.RequestException as e:
        print(f"  [ERR]  {filename} ({e})")
        return False


def create_word_document(
    cards: list[tuple[str, int]],
    save_dir: str,
    deck_name: str,
    cols: int = 2,
    rows: int = 4,
):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)

    image_paths = []
    for card_id, qty in cards:
        filepath = os.path.join(save_dir, f"{card_id}.png")
        if os.path.exists(filepath):
            for _ in range(qty):
                image_paths.append(filepath)

    rotated_dir = os.path.join(save_dir, "_rotated")
    os.makedirs(rotated_dir, exist_ok=True)
    rotated_cache = {}
    for img_path in set(image_paths):
        rotated_path = os.path.join(rotated_dir, os.path.basename(img_path))
        with Image.open(img_path) as img:
            rotated = img.rotate(90, expand=True)
            rotated.save(rotated_path, dpi=(DPI, DPI))
        rotated_cache[img_path] = rotated_path

    COLS = cols
    ROWS = rows
    CARDS_PER_PAGE = COLS * ROWS
    img_w = Mm(88)
    img_h = Mm(63)

    for page_start in range(0, len(image_paths), CARDS_PER_PAGE):
        if page_start > 0:
            doc.add_page_break()

        page_cards = image_paths[page_start : page_start + CARDS_PER_PAGE]
        table = doc.add_table(rows=ROWS, cols=COLS)
        table.autofit = False

        for col in table.columns:
            for cell in col.cells:
                cell.width = img_w

        for i, img_path in enumerate(page_cards):
            row_idx = i // COLS
            col_idx = i % COLS
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].alignment = 1

            run = cell.paragraphs[0].add_run()
            run.add_picture(rotated_cache[img_path], width=img_w, height=img_h)

    docx_path = os.path.join(save_dir, f"{deck_name}.docx")
    doc.save(docx_path)
    print(f"Word document saved to {docx_path}")


def main():
    use_cached = "--use-cached" in sys.argv
    args = [a for a in sys.argv[1:] if a != "--use-cached"]

    if len(args) < 1:
        print("Usage: python digimoncrawler.py [--use-cached] <recipe_url>")
        print()
        print("Example:")
        print(
            '  python digimoncrawler.py "https://digimon-cg-guide.com/recipe-creater/?recipe=bt14-0014_bt22-0083&deckname=MyDeck"'
        )
        sys.exit(1)

    url = args[0]
    cards, deck_name = parse_recipe_url(url)

    safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in deck_name)
    save_dir = os.path.join("digimon_cards", safe_name)
    os.makedirs(save_dir, exist_ok=True)

    unique_cards = {}
    for card_id, qty in cards:
        unique_cards[card_id] = qty

    print(f"Deck: {deck_name}")
    print(f"Cards: {len(cards)} entries, {len(unique_cards)} unique")
    print(f"Save to: {save_dir}/")
    print()

    if use_cached:
        print("Using cached images.")
    else:
        success = 0
        fail = 0
        for card_id, qty in unique_cards.items():
            ok = download_card_image(card_id, save_dir)
            if ok:
                success += 1
            else:
                fail += 1
            time.sleep(0.3)

        print()
        print(f"Done! {success} downloaded, {fail} failed.")

    summary_path = os.path.join(save_dir, "decklist.txt")
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(f"Deck: {deck_name}\n")
        f.write(f"Source: {url}\n\n")
        for card_id, qty in cards:
            f.write(f"{qty}x {card_id}\n")
    print(f"Deck list saved to {summary_path}")

    create_word_document(cards, save_dir, safe_name)
    create_word_document(cards, save_dir, f"{safe_name}_3x3", cols=3, rows=3)


if __name__ == "__main__":
    main()
